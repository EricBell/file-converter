"""
DOCX writer implementation using python-docx.
"""

from pathlib import Path
from typing import Union, Optional
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import BaseWriter
from ..core.document import Document, ElementType
from ..core.footer import FooterConfig
from ..core.lockfile import cleanup_lock_files


class DocxWriter(BaseWriter):
    """Writer for DOCX files using python-docx."""

    def write(self, document: Document, output_path: Union[str, Path],
              footer_config: Optional[FooterConfig] = None) -> None:
        """Write a document to a DOCX file."""
        output_path = Path(output_path)

        if not self.supports_format(output_path):
            raise ValueError(f"Unsupported output format: "
                           f"{output_path.suffix}")

        try:
            docx_doc = self._create_docx_document(document, footer_config)
            docx_doc.save(output_path)
        except Exception as e:
            raise IOError(f"Error writing to file {output_path}: {str(e)}")
        finally:
            # Always clean up lock files, even if an error occurred
            cleanup_lock_files(output_path)

    def to_string(self, document: Document) -> str:
        """
        Convert a document to string representation.

        Note: DOCX format doesn't have a meaningful string representation.
        This returns a summary of the document structure.
        """
        if not isinstance(document, Document):
            raise ValueError("Input must be a Document object")

        parts = []
        if document.title:
            parts.append(f"Title: {document.title}")

        parts.append(f"Elements: {len(document.elements)}")

        element_counts = {}
        for element in document.elements:
            element_type = element.element_type.value
            element_counts[element_type] = element_counts.get(element_type, 0) + 1

        for element_type, count in element_counts.items():
            parts.append(f"  {element_type}: {count}")

        return "\n".join(parts)

    def _create_docx_document(self, document: Document,
                              footer_config: Optional[FooterConfig] = None) -> DocxDocument:
        """Create a python-docx Document from our Document object."""
        docx_doc = DocxDocument()

        # Add title if present
        if document.title:
            title = docx_doc.add_heading(document.title, level=0)

        # Process each element
        for element in document.elements:
            self._add_element_to_docx(element, docx_doc)

        # Add footer if configured
        if footer_config and footer_config.enabled:
            self._add_footer(docx_doc, footer_config)

        return docx_doc

    def _add_footer(self, docx_doc: DocxDocument, footer_config: FooterConfig) -> None:
        """
        Add footer to the document.

        Args:
            docx_doc: The python-docx Document object
            footer_config: Footer configuration
        """
        section = docx_doc.sections[0]

        # Configure for different odd/even pages if double-sided
        if footer_config.layout == "double":
            section.different_first_page_header_footer = False
            section.different_odd_even_pages_header_footer = True

            # Add odd page footer (first page footer becomes odd page footer)
            odd_footer = section.footer
            self._format_footer(odd_footer, footer_config, page_number=1)

            # Add even page footer
            even_footer = section.even_page_footer
            self._format_footer(even_footer, footer_config, page_number=2)
        else:
            # Single-sided: same footer on all pages
            footer = section.footer
            self._format_footer(footer, footer_config, page_number=1)

    def _format_footer(self, footer, footer_config: FooterConfig, page_number: int) -> None:
        """
        Format a footer with left and right aligned text.

        Args:
            footer: The footer object to format
            footer_config: Footer configuration
            page_number: Page number to determine layout (for double-sided)
        """
        # Get the footer text for this page type
        left_text, right_text = footer_config.get_footer_for_page(page_number)

        # Clear any existing content
        for paragraph in footer.paragraphs:
            paragraph.clear()

        # Add paragraph with tab stops for alignment
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Add left-aligned text
        run_left = para.add_run(left_text)
        run_left.font.size = Pt(9)

        # Add tab and right-aligned text
        # Note: We use a simple approach with tabs. For page numbers in DOCX,
        # we'll use the placeholder text since actual page numbers would require
        # field codes which are more complex.
        para.add_run('\t' + right_text)
        para.runs[1].font.size = Pt(9)

        # Set tab stop for right alignment
        from docx.shared import RGBColor
        para.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)

    def _add_element_to_docx(self, element, docx_doc: DocxDocument) -> None:
        """Add a document element to the DOCX document."""
        if element.element_type == ElementType.HEADING:
            self._add_heading_to_docx(element, docx_doc)
        elif element.element_type == ElementType.PARAGRAPH:
            self._add_paragraph_to_docx(element, docx_doc)
        elif element.element_type == ElementType.LIST:
            self._add_list_to_docx(element, docx_doc)
        elif element.element_type == ElementType.CODE_BLOCK:
            self._add_code_block_to_docx(element, docx_doc)
        else:
            # Fallback: treat as paragraph
            if element.content:
                docx_doc.add_paragraph(element.content)

    def _add_heading_to_docx(self, heading, docx_doc: DocxDocument) -> None:
        """Add a heading element to DOCX."""
        level = heading.attributes.get('level', 1)
        # Word headings are limited to levels 1-9
        level = min(max(level, 1), 9)
        docx_doc.add_heading(heading.content, level=level)

    def _add_paragraph_to_docx(self, paragraph, docx_doc: DocxDocument) -> None:
        """Add a paragraph element to DOCX."""
        docx_doc.add_paragraph(paragraph.content)

    def _add_list_to_docx(self, list_element, docx_doc: DocxDocument, indent_level=0) -> None:
        """
        Add a list element to DOCX with support for nested lists.

        Args:
            list_element: The DocumentList to convert
            docx_doc: The python-docx Document object
            indent_level: Current nesting level (0 = top level)
        """
        if not hasattr(list_element, 'items') or not list_element.items:
            return

        # python-docx doesn't have built-in list support, so we'll simulate it
        # with indented paragraphs
        is_ordered = list_element.attributes.get('ordered', False)

        # Calculate indentation based on nesting level
        base_indent = 0.25 + (indent_level * 0.5)  # Increase indent for each level

        for i, item in enumerate(list_element.items, 1):
            if is_ordered:
                bullet = f"{i}."
            else:
                bullet = "•"

            para = docx_doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(base_indent)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            run = para.add_run(f"{bullet} {item.content}")

            # Process nested children if any
            if hasattr(item, 'children') and item.children:
                for child_list in item.children:
                    self._add_list_to_docx(child_list, docx_doc, indent_level + 1)

    def _add_code_block_to_docx(self, code_block,
                               docx_doc: DocxDocument) -> None:
        """Add a code block element to DOCX."""
        # Add code as a paragraph with monospace font
        para = docx_doc.add_paragraph()
        run = para.add_run(code_block.content)
        run.font.name = 'Courier New'
        # Use a safe default size instead of accessing styles
        try:
            run.font.size = docx_doc.styles['Normal'].font.size
        except (KeyError, TypeError):
            # Fallback if styles access fails (e.g., in tests)
            pass

        # Add some styling to make it look like a code block
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Inches(0.1)
        para.paragraph_format.space_after = Inches(0.1)

    def supports_format(self, file_path: Union[str, Path]) -> bool:
        """Check if this writer supports the given file format."""
        return Path(file_path).suffix.lower() == '.docx'

    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        """Get list of supported file extensions."""
        return ['.docx']